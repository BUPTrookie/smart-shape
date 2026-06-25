from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips


BASE_FONT = "Calibri"
CJK_FONT = "Microsoft YaHei"
TITLE_FONT = "Microsoft YaHei"
CODE_FONT = "Consolas"


def find_source() -> Path:
    docs = Path.cwd() / "docs"
    matches = sorted(docs.glob("*结题报告.md"))
    if not matches:
        raise FileNotFoundError("未找到结题报告 Markdown 源文件")
    return matches[0]


def set_font(run, *, name=BASE_FONT, size=None, bold=None, italic=None, color=None):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), CJK_FONT if name != CODE_FONT else name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_format(
    paragraph,
    *,
    before=0,
    after=6,
    line=1.1,
    align=WD_ALIGN_PARAGRAPH.LEFT,
    left=0,
    right=0,
    first=0,
):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    fmt.left_indent = Inches(left)
    fmt.right_indent = Inches(right)
    fmt.first_line_indent = Inches(first)
    paragraph.alignment = align


def set_run_text(paragraph, text, *, size=11, bold=False, italic=False, color=None):
    token_re = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*|\*[^*]+\*)")
    for token in token_re.split(text):
        if not token:
            continue
        if token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_font(run, name=CODE_FONT, size=size, bold=bold, italic=italic, color=color)
        elif token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_font(run, name=BASE_FONT, size=size, bold=True, italic=italic, color=color)
        elif token.startswith("*") and token.endswith("*"):
            run = paragraph.add_run(token[1:-1])
            set_font(run, name=BASE_FONT, size=size, bold=bold, italic=True, color=color)
        else:
            run = paragraph.add_run(token)
            set_font(run, name=BASE_FONT, size=size, bold=bold, italic=italic, color=color)


def is_table_separator(line: str) -> bool:
    stripped = line.strip()
    if "|" not in stripped:
        return False
    cells = [c.strip() for c in stripped.strip("|").split("|")]
    return len(cells) >= 2 and all(re.fullmatch(r"[:\- ]+", cell or "") for cell in cells)


def split_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def weight_text(value: str) -> int:
    total = 0
    for ch in re.sub(r"`([^`]+)`", r"\1", value):
        total += 2 if ord(ch) > 127 else 1
    return max(total, 3)


def calc_column_widths(rows: list[list[str]], total_width=9360) -> list[int]:
    scores = [0] * len(rows[0])
    for row in rows:
        for idx, cell in enumerate(row):
            scores[idx] = max(scores[idx], weight_text(cell))
    widths = [int(round(total_width * score / sum(scores))) for score in scores]
    widths[-1] += total_width - sum(widths)
    return widths


def set_table_geometry(table, widths: list[int], total_width=9360):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    def ensure(tag):
        node = tbl_pr.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            tbl_pr.append(node)
        return node

    tbl_w = ensure("w:tblW")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(total_width))

    tbl_ind = ensure("w:tblInd")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "120")

    layout = ensure("w:tblLayout")
    layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for col_idx, width in enumerate(widths):
        table.columns[col_idx].width = Twips(width)

    for row in table.rows:
        for col_idx, cell in enumerate(row.cells):
            cell.width = Twips(widths[col_idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths[col_idx]))

            tc_mar = tc_pr.find(qn("w:tcMar"))
            if tc_mar is None:
                tc_mar = OxmlElement("w:tcMar")
                tc_pr.append(tc_mar)
            for side, width in {"top": 80, "bottom": 80, "start": 120, "end": 120}.items():
                node = tc_mar.find(qn(f"w:{side}"))
                if node is None:
                    node = OxmlElement(f"w:{side}")
                    tc_mar.append(node)
                node.set(qn("w:w"), str(width))
                node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")


def make_table(doc: Document, lines: list[str]):
    rows = [split_table_row(lines[0])]
    for line in lines[2:]:
        rows.append(split_table_row(line))
    widths = calc_column_widths(rows)
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    set_table_geometry(table, widths)

    for r_idx, row in enumerate(rows):
        for c_idx, text in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.text = ""
            para = cell.paragraphs[0]
            set_paragraph_format(para, before=0, after=0, line=1.0)
            if r_idx == 0:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_run_text(para, text, size=10, bold=(r_idx == 0))
            if r_idx == 0:
                shade_cell(cell, "E8EEF5")
    return table


def make_code_block(doc: Document, lines: list[str]):
    para = doc.add_paragraph()
    set_paragraph_format(para, before=2, after=6, line=1.0, left=0.15)
    ppr = para._p.get_or_add_pPr()
    shd = ppr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        ppr.append(shd)
    shd.set(qn("w:fill"), "F7F7F7")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    for idx, line in enumerate(lines):
        run = para.add_run(line)
        set_font(run, name=CODE_FONT, size=9.5)
        if idx != len(lines) - 1:
            run.add_break(WD_BREAK.LINE)


def make_caption(doc: Document, text: str):
    para = doc.add_paragraph()
    set_paragraph_format(para, before=2, after=4, line=1.0)
    set_run_text(para, text, size=10, italic=True, color="555555")


def parse_blocks(text: str):
    lines = text.splitlines()
    i = 0

    if i < len(lines) and lines[i].startswith("# "):
        yield ("title", lines[i][2:].strip())
        i += 1
    if i < len(lines) and lines[i].startswith("## "):
        yield ("subtitle", lines[i][3:].strip())
        i += 1

    while i < len(lines) and lines[i].strip().startswith("**"):
        yield ("meta", lines[i].strip())
        i += 1

    while i < len(lines) and not lines[i].strip():
        i += 1
    if i < len(lines) and lines[i].strip() == "---":
        i += 1

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if not stripped or stripped == "---":
            yield ("blank", "")
            i += 1
            continue

        if stripped.startswith("```"):
            i += 1
            code = []
            while i < len(lines) and lines[i].strip() != "```":
                code.append(lines[i])
                i += 1
            if i < len(lines):
                i += 1
            yield ("code", code)
            continue

        if i + 1 < len(lines) and is_table_separator(lines[i + 1]) and "|" in line:
            table_lines = [line, lines[i + 1]]
            i += 2
            while i < len(lines) and lines[i].strip() and "|" in lines[i]:
                table_lines.append(lines[i])
                i += 1
            yield ("table", table_lines)
            continue

        if stripped.startswith(">"):
            quote_lines = [stripped[1:].strip()]
            i += 1
            while i < len(lines) and lines[i].strip().startswith(">"):
                quote_lines.append(lines[i].strip()[1:].strip())
                i += 1
            yield ("caption", " ".join(quote_lines))
            continue

        heading = re.match(r"^(#{2,4})\s+(.*)$", stripped)
        if heading:
            yield ("heading", (len(heading.group(1)) - 1, heading.group(2).strip()))
            i += 1
            continue

        bullet = re.match(r"^[-*+]\s+(.*)$", stripped)
        if bullet:
            yield ("bullet", bullet.group(1))
            i += 1
            continue

        numbered = re.match(r"^\d+\.\s+(.*)$", stripped)
        if numbered:
            yield ("number", numbered.group(1))
            i += 1
            continue

        para_lines = [stripped]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if not nxt:
                break
            if nxt == "---" or nxt.startswith("#") or nxt.startswith(">") or nxt.startswith("```"):
                break
            if re.match(r"^[-*+]\s+", nxt) or re.match(r"^\d+\.\s+", nxt):
                break
            if i + 1 < len(lines) and is_table_separator(lines[i + 1]) and "|" in lines[i]:
                break
            para_lines.append(nxt)
            i += 1
        yield ("para", " ".join(para_lines))


def build_doc():
    src = find_source()
    text = src.read_text(encoding="utf-8-sig")

    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = BASE_FONT
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:ascii"), BASE_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), BASE_FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = doc.styles[style_name]
        style.font.name = BASE_FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:ascii"), BASE_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), BASE_FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.1

    if "List Bullet" in doc.styles:
        lb = doc.styles["List Bullet"]
        lb.font.name = BASE_FONT
        lb.font.size = Pt(11)
        lb._element.rPr.rFonts.set(qn("w:ascii"), BASE_FONT)
        lb._element.rPr.rFonts.set(qn("w:hAnsi"), BASE_FONT)
        lb._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
        lb.paragraph_format.space_after = Pt(4)
        lb.paragraph_format.line_spacing = 1.1
    if "List Number" in doc.styles:
        ln = doc.styles["List Number"]
        ln.font.name = BASE_FONT
        ln.font.size = Pt(11)
        ln._element.rPr.rFonts.set(qn("w:ascii"), BASE_FONT)
        ln._element.rPr.rFonts.set(qn("w:hAnsi"), BASE_FONT)
        ln._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
        ln.paragraph_format.space_after = Pt(4)
        ln.paragraph_format.line_spacing = 1.1

    for kind, value in parse_blocks(text):
        if kind == "title":
            para = doc.add_paragraph()
            set_paragraph_format(para, before=0, after=3, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
            set_run_text(para, value, size=20, bold=True, color="000000")
        elif kind == "subtitle":
            para = doc.add_paragraph()
            set_paragraph_format(para, before=0, after=4, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
            set_run_text(para, value, size=14, bold=True, color="000000")
        elif kind == "meta":
            para = doc.add_paragraph()
            set_paragraph_format(para, before=0, after=2, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
            cleaned = re.sub(r"\*\*(.*?)\*\*", r"\1", value)
            cleaned = cleaned.strip("*")
            set_run_text(para, cleaned, size=10.5, color="444444")
        elif kind == "blank":
            doc.add_paragraph().paragraph_format.space_after = Pt(0)
        elif kind == "heading":
            level, heading_text = value
            para = doc.add_paragraph(style=f"Heading {level}")
            set_paragraph_format(para, before=16 if level == 1 else 12 if level == 2 else 8, after=8 if level == 1 else 6 if level == 2 else 4, line=1.1)
            para.paragraph_format.keep_with_next = True
            set_run_text(para, heading_text, size=16 if level == 1 else 13 if level == 2 else 12, bold=True)
        elif kind == "para":
            para = doc.add_paragraph()
            set_paragraph_format(para, before=0, after=6, line=1.1)
            set_run_text(para, value, size=11)
        elif kind == "bullet":
            para = doc.add_paragraph(style="List Bullet")
            set_paragraph_format(para, before=0, after=4, line=1.1)
            set_run_text(para, value, size=11)
        elif kind == "number":
            para = doc.add_paragraph(style="List Number")
            set_paragraph_format(para, before=0, after=4, line=1.1)
            set_run_text(para, value, size=11)
        elif kind == "caption":
            make_caption(doc, value)
        elif kind == "code":
            make_code_block(doc, value)
        elif kind == "table":
            make_table(doc, value)

    out = src.with_suffix(".docx")
    doc.save(out)
    return out


if __name__ == "__main__":
    out = build_doc()
    print(out)
